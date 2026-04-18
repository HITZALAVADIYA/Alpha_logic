from flask import Flask, render_template, request, jsonify, redirect, url_for
from flask_cors import CORS
import json
import re
import traceback
import pdfplumber
import docx
import os
import io 
import psycopg2 
from psycopg2.extras import RealDictCursor
from supabase import create_client, Client
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from dotenv import load_dotenv
from datetime import datetime

# Load environment variables
load_dotenv()

# Orchestrator import
try:
    from orchestrator import run_sentinel_analysis 
except ImportError:
    print("❌ CRITICAL: orchestrator.py missing in folder!")

app = Flask(__name__)
CORS(app)
app.secret_key = os.environ.get("SENTINEL_SECRET_KEY", "sentinel-super-secret-2026")

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "signin"

print("\n" + "="*60)
print("🚀 RESUME SENTINEL - INITIALIZATION")
print("="*60)

# ==========================================
# 1️⃣ CLOUD DATABASE (SUPABASE) - PRIMARY ⭐
# ==========================================
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://fdrexxunukdbdyathdiv.supabase.co")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "sb_publishable_pM5tfO7H9RsCCAOjNcF-XA_HjJkpZqn")

SUPABASE_CONNECTED = False
supabase = None

try:
    if SUPABASE_URL and SUPABASE_KEY:
        supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
        # Test connection
        response = supabase.table('app_users').select("count", count="exact").limit(1).execute()
        SUPABASE_CONNECTED = True
        print("☁️  [SUPABASE] ✅ Connected successfully!")
    else:
        print("⚠️  [SUPABASE] Credentials not found")
except Exception as e:
    print(f"⚠️  [SUPABASE] Connection failed: {str(e)[:80]}")
    SUPABASE_CONNECTED = False

# ==========================================
# 2️⃣ LOCAL DATABASE (POSTGRESQL) - OPTIONAL
# ==========================================
DB_HOST = os.environ.get("DB_HOST", "localhost")
DB_PORT = os.environ.get("DB_PORT", "5432")
DB_NAME = os.environ.get("DB_NAME", "resume_db_schema")
DB_USER = os.environ.get("DB_USER", "postgres")
DB_PASS = os.environ.get("DB_PASSWORD", "root")

POSTGRESQL_CONNECTED = False

def get_db_connection():
    """Get PostgreSQL connection"""
    try:
        conn = psycopg2.connect(
            host=DB_HOST,
            port=DB_PORT,
            database=DB_NAME,
            user=DB_USER,
            password=DB_PASS
        )
        return conn
    except Exception as e:
        return None

# Test PostgreSQL connection
test_conn = get_db_connection()
if test_conn:
    POSTGRESQL_CONNECTED = True
    print("🐘 [POSTGRESQL] ✅ Connected successfully!")
    test_conn.close()
else:
    print("⚠️  [POSTGRESQL] Not available (optional)")

# ==========================================
# DATABASE SCHEMA INITIALIZATION
# ==========================================
def init_postgresql_database():
    """Initialize PostgreSQL tables (optional)"""
    if not POSTGRESQL_CONNECTED:
        return False
    
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # Profiles table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS profiles (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(255) NOT NULL,
                    resume_text TEXT,
                    jd_text TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Scan reports table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS scan_reports (
                    id SERIAL PRIMARY KEY,
                    candidate_name VARCHAR(255),
                    match_score INTEGER,
                    status VARCHAR(255),
                    ai_analysis JSONB,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # App users table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS app_users (
                    id SERIAL PRIMARY KEY,
                    full_name VARCHAR(255),
                    username VARCHAR(255) UNIQUE,
                    email VARCHAR(255) UNIQUE NOT NULL,
                    password VARCHAR(255) NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Insert default admin
            cursor.execute("""
                INSERT INTO app_users (full_name, username, email, password)
                VALUES ('System Admin', 'admin', 'admin@sentinel.ai', 'admin123')
                ON CONFLICT (email) DO NOTHING
            """)
            
            conn.commit()
            cursor.close()
            conn.close()
            print("✅ [POSTGRESQL] Tables initialized!")
            return True
    except Exception as e:
        print(f"⚠️  [POSTGRESQL] Schema init failed: {str(e)[:80]}")
        return False

if POSTGRESQL_CONNECTED:
    init_postgresql_database()

print("="*60 + "\n")

# ==========================================
# AUTHENTICATION MODELS
# ==========================================
class User(UserMixin):
    def __init__(self, user_id, email, full_name="User"): 
        self.id = str(user_id)
        self.email = email
        self.full_name = full_name

@login_manager.user_loader
def load_user(user_id):
    """Load user from Supabase (primary) or PostgreSQL (fallback)"""
    if SUPABASE_CONNECTED:
        try:
            response = supabase.table('app_users').select("id, email, full_name").eq('id', int(user_id)).execute()
            if response.data and len(response.data) > 0:
                user = response.data[0]
                return User(user['id'], user['email'], user.get('full_name', 'User'))
        except:
            pass
    
    if POSTGRESQL_CONNECTED:
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT id, email, full_name FROM app_users WHERE id = %s", (user_id,))
                row = cursor.fetchone()
                cursor.close()
                conn.close()
                return User(row[0], row[1], row[2]) if row else None
            except:
                return None
    
    return None

@login_manager.unauthorized_handler
def unauthorized():
    """Handle unauthorized access"""
    if request.path.startswith("/api/"):
        return jsonify({"login_required": True}), 401
    return redirect(url_for("signin", next=request.path))

# ==========================================
# 📄 FILE EXTRACTION UTILITIES
# ==========================================
def extract_text_from_pdf(file_stream):
    """Extract text from PDF"""
    try:
        pdf = pdfplumber.open(file_stream)
        text = ""
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        pdf.close()
        return text.strip()
    except Exception as e:
        print(f"❌ PDF extraction error: {str(e)[:80]}")
        return ""

def extract_text_from_docx(file_stream):
    """Extract text from DOCX with proper error handling"""
    try:
        doc = docx.Document(file_stream)
        text = []
        for para in doc.paragraphs:
            if para.text.strip(): text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data: text.append(" | ".join(row_data))
        return "\n".join(text).strip()
    except Exception as e:
        print(f"❌ DOCX extraction error: {str(e)[:80]}")
        return ""

def extract_text_from_file(file):
    """Universal file extractor"""
    filename = file.filename.lower()
    try:
        file.seek(0)
        content = file.read()
        if not content: return ""
        if filename.endswith('.pdf'): return extract_text_from_pdf(io.BytesIO(content))
        elif filename.endswith('.docx'): return extract_text_from_docx(io.BytesIO(content))
        elif filename.endswith('.txt'): return content.decode('utf-8').strip()
        return ""
    except Exception as e:
        print(f"❌ File extraction error: {str(e)[:80]}")
        return ""

# ==========================================
# DATABASE SAVE FUNCTIONS
# ==========================================

def save_report_to_supabase(candidate_name, match_score, status, ai_analysis):
    if not SUPABASE_CONNECTED: return False
    try:
        supabase.table('scan_reports').insert({
            "candidate_name": candidate_name,
            "match_score": match_score,
            "status": status,
            "ai_analysis": ai_analysis
        }).execute()
        return True
    except: return False

def save_report_to_postgresql(candidate_name, match_score, status, ai_analysis):
    if not POSTGRESQL_CONNECTED: return False
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO scan_reports (candidate_name, match_score, status, ai_analysis)
                VALUES (%s, %s, %s, %s)
            """, (candidate_name, match_score, status, json.dumps(ai_analysis)))
            conn.commit()
            cursor.close()
            conn.close()
            return True
    except: return False

def get_user_from_supabase(email):
    if not SUPABASE_CONNECTED: return None
    try:
        response = supabase.table('app_users').select("*").eq('email', email).execute()
        return response.data[0] if response.data else None
    except: return None

def get_user_from_postgresql(email):
    if not POSTGRESQL_CONNECTED: return None
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            cursor.execute("SELECT * FROM app_users WHERE email = %s", (email,))
            user = cursor.fetchone()
            cursor.close()
            conn.close()
            return dict(user) if user else None
        except: return None
    return None

# ==========================================
# 🚀 CORE ROUTES (BULK UPLOAD SUPPORT ADDED)
# ==========================================

@app.route('/')
def home():
    return render_template('index.html', 
                         logged_in=current_user.is_authenticated,
                         supabase_connected=SUPABASE_CONNECTED,
                         postgresql_connected=POSTGRESQL_CONNECTED)

@app.route('/signin', methods=['GET', 'POST'])
def signin():
    if current_user.is_authenticated: return redirect(url_for('home'))
    error = None
    email = request.form.get('email', '').strip() if request.method == 'POST' else ''
    if request.method == 'POST':
        password = request.form.get('password', '')
        user = get_user_from_supabase(email) or get_user_from_postgresql(email)
        if user and user.get('password') == password:
            login_user(User(user['id'], user['email'], user.get('full_name', 'User')))
            return redirect(request.args.get('next') or url_for('home'))
        error = "❌ Invalid email or password"
    return render_template('signin.html', error=error, email=email, panel='signin')

@app.route('/register', methods=['POST'])
def register():
    full_name = request.form.get('full_name', '').strip()
    email = request.form.get('email', '').strip()
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')
    if not all([full_name, email, username, password]):
        return render_template('signin.html', reg_error="All fields are required", panel='register')
    
    # Supabase Register
    if SUPABASE_CONNECTED:
        try:
            supabase.table('app_users').insert({"full_name": full_name, "username": username, "email": email, "password": password}).execute()
            return render_template('signin.html', registered=True, panel='signin')
        except: pass
    
    # Postgres Register Fallback
    if POSTGRESQL_CONNECTED:
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("INSERT INTO app_users (full_name, username, email, password) VALUES (%s, %s, %s, %s)", (full_name, username, email, password))
            conn.commit()
            cursor.close()
            conn.close()
            return render_template('signin.html', registered=True, panel='signin')
        except: pass
    return render_template('signin.html', reg_error="Database unavailable", panel='register')

@app.route('/signout')
def signout():
    if current_user.is_authenticated: logout_user()
    return render_template('sigout.html')

# 🔥 NAYA ROUTE: SUPPORT FOR BULK AND SINGLE ANALYSIS
@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    """Support for Single and Bulk Uploads"""
    try:
        print("\n📥 Processing Analysis Batch...")
        jd_text = request.form.get('jd', '').strip()
        resumes_to_process = []

        # 1. Gather files from 'resume_files' (Bulk Input)
        files = request.files.getlist('resume_files')
        
        # 2. Backward Compatibility for single 'resume_file'
        if not files or (len(files) == 1 and not files[0].filename):
            single = request.files.get('resume_file')
            if single and single.filename:
                files = [single]

        # 3. Extract text from each file
        for file in files:
            if file.filename:
                text = extract_text_from_file(file)
                if len(text.strip()) > 20:
                    resumes_to_process.append({"name": file.filename, "text": text})

        # 4. Fallback to manual textarea if no files
        if not resumes_to_process:
            manual_text = request.form.get('resume', '').strip()
            if manual_text:
                resumes_to_process.append({"name": request.form.get('candidate_name', 'Manual Entry'), "text": manual_text})

        if not resumes_to_process: raise ValueError("No valid resumes found.")
        if len(jd_text) < 50: raise ValueError("Job description too short.")

        print(f"📊 Batch Size: {len(resumes_to_process)} resumes")

        batch_results = []

        # 5. Process in a Loop
        for item in resumes_to_process:
            print(f"🤖 Processing: {item['name']}")
            
            # 🔥 HACKATHON FIX: Truncate texts to strictly prevent TPM Rate Limit on Groq Free Tier
            safe_resume_text = item['text'][:1800] # Max ~450 tokens
            safe_jd_text = jd_text[:1000] # Max ~250 tokens
            
            raw_output = run_sentinel_analysis(safe_resume_text, safe_jd_text)
            
            json_match = re.search(r'\{.*\}', str(raw_output), re.DOTALL)
            if json_match:
                data = json.loads(json_match.group())
                report = {
                    "filename": item['name'],
                    "score": data.get('score', 0),
                    "skills": data.get('skills', []),
                    "insight": data.get('insight', "Analysis complete"),
                    "questions": data.get('questions', []),
                    "cover_letter": data.get('cover_letter', "Could not generate cover letter. Pipeline execution ran too fast or encountered an issue."),
                    "cold_email": data.get('cold_email', "Could not generate cold email. Pipeline execution ran too fast or encountered an issue.")
                }
                
                # Save each one to DBs
                status = "APPROVED" if report['score'] >= 80 else "REVIEW_NEEDED"
                save_report_to_supabase(report['filename'], report['score'], status, report)
                save_report_to_postgresql(report['filename'], report['score'], status, report)
                
                batch_results.append(report)
                
            import time
            time.sleep(4) # Rate limit bypass for Groq

        print(f"✅ Batch processed successfully. Processed {len(batch_results)} candidates.")
        
        # Return list for bulk, or single object for legacy frontend
        return jsonify(batch_results if len(batch_results) > 1 else batch_results[0]), 200
        
    except Exception as e:
        print(f"❌ Batch error: {str(e)}")
        return jsonify({"score": 0, "insight": f"Analysis failed: {str(e)}"}), 400

@app.route('/profiles', methods=['GET', 'POST'])
@login_required
def manage_profiles():
    if SUPABASE_CONNECTED:
        try:
            if request.method == 'GET':
                res = supabase.table('profiles').select("id, name, created_at").order('created_at', desc=True).execute()
                return jsonify({"profiles": res.data if res.data else []})
            if request.method == 'POST':
                data = request.get_json()
                res = supabase.table('profiles').insert({"name": data.get('name'), "resume_text": data.get('resume_text'), "jd_text": data.get('jd_text')}).execute()
                return jsonify({"id": res.data[0]['id'], "message": "✅ Saved!"}), 201
        except: pass
    
    if POSTGRESQL_CONNECTED:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        if request.method == 'GET':
            cursor.execute("SELECT id, name, created_at FROM profiles ORDER BY created_at DESC")
            profiles = cursor.fetchall()
            cursor.close(); conn.close()
            return jsonify({"profiles": [dict(p) for p in profiles]})
        if request.method == 'POST':
            data = request.get_json()
            cursor.execute("INSERT INTO profiles (name, resume_text, jd_text) VALUES (%s,%s,%s) RETURNING id", (data.get('name'), data.get('resume_text'), data.get('jd_text')))
            new_id = cursor.fetchone()['id']; conn.commit()
            cursor.close(); conn.close()
            return jsonify({"id": new_id, "message": "✅ Saved!"}), 201
    return jsonify({"error": "No database available"}), 503

@app.route('/reports')
@login_required
def get_reports():
    try:
        if SUPABASE_CONNECTED:
            res = supabase.table('scan_reports').select("*").order('created_at', desc=True).limit(100).execute()
            return jsonify({"status": "success", "data": res.data if res.data else []})
        if POSTGRESQL_CONNECTED:
            conn = get_db_connection()
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            cursor.execute("SELECT * FROM scan_reports ORDER BY created_at DESC LIMIT 100")
            rows = cursor.fetchall(); cursor.close(); conn.close()
            return jsonify({"status": "success", "data": [dict(r) for r in rows]})
    except: pass
    return jsonify({"status": "error", "message": "Database error"}), 500

@app.route('/api/chat', methods=['POST'])
@login_required
def api_chat():
    try:
        user_msg = request.json.get('message', '')
        api_key = os.environ.get("GROQ_API_KEY", "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")
        import requests
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "llama-3.1-8b-instant",
            "messages": [
                {"role": "system", "content": "You are a tough Technical Interviewer. Keep answers brief (max 2 sentences). Ask follow up questions based on the candidate's answer."},
                {"role": "user", "content": user_msg}
            ]
        }
        r = requests.post("https://api.groq.com/openai/v1/chat/completions", headers=headers, json=payload).json()
        if "choices" in r:
            return jsonify({"reply": r["choices"][0]["message"]["content"]})
        else:
            return jsonify({"reply": f"API Error: {r.get('error', r)}"})
    except Exception as e:
        return jsonify({"reply": f"Mock API error: {str(e)}"})

@app.route('/status')
def status():
    return jsonify({
        "supabase": "✅ Connected" if SUPABASE_CONNECTED else "❌ Disconnected",
        "postgresql": "✅ Connected" if POSTGRESQL_CONNECTED else "❌ Disconnected"
    })

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 RESUME SENTINEL - STARTING SERVER")
    print("📍 Server running on: http://localhost:5000")
    print("="*60 + "\n")
    app.run(debug=True, port=5000, host='0.0.0.0')